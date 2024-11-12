from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK
from scripts.data_fetcher import fetch_data
from datetime import date

def get_polish_month(month):
    """
    Funkcja tlumaczaca miesiac na polski.
    """
    months = [
        "styczeń", "luty", "marzec", "kwiecień", "maj", "czerwiec", 
        "lipiec", "sierpień", "wrzesień", "październik", "listopad", "grudzień"
    ]
    return months[month - 1]

def create_document():
    """
    Tworzy obiekt dokumentu i ustawia marginesy, styl tekstu i naglowki.
    """
    doc = Document()

    # Ustawienie marginesow na 2 cm
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Definiowanie stylow dokumentu
    define_styles(doc)
    
    # Dodanie naglowka
    add_header(doc)

    return doc

def define_styles(doc):
    """
    Ustawia style dla tekstu, tytulu, tekstu pogrubionego i tekstu wcietego.
    """
    # Styl ogolny
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.0
    normal_style.paragraph_format.space_after = 0

    # Styl tytulu
    title_style = doc.styles.add_style('MyTitle', 1)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.paragraph_format.line_spacing = 1.0
    title_style.paragraph_format.space_after = 0

    # Styl pogrubiony
    bold_style = doc.styles.add_style('NormalBold', 1)
    bold_style.font.name = 'Times New Roman'
    bold_style.font.size = Pt(12)
    bold_style.font.bold = True
    bold_style.paragraph_format.line_spacing = 1.0
    bold_style.paragraph_format.space_after = 0

    # Styl z wcieciem
    tab_indent_style = doc.styles.add_style('TabulatorStyl', 1)
    tab_indent_style.font.name = 'Times New Roman'
    tab_indent_style.font.size = Pt(12)
    tab_indent_style.paragraph_format.left_indent = Pt(40)
    tab_indent_style.paragraph_format.line_spacing = 1.0
    tab_indent_style.paragraph_format.space_after = 0

def add_header(doc):
    """
    Dodaje naglowek do dokumentu.
    """
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "Analiza Komputerow Emisyjnych"
        paragraph.alignment = 2
        paragraph.style.font.name = 'Times New Roman'
        paragraph.style.font.size = Pt(12)

def add_document_title(doc, date):
    """
    Dodaje tytul dokumentu oraz naglowek z data.
    """
    doc.add_paragraph(f'Ostrzeszow, dn. {date} r.\n\n', style='Normal').alignment = 2
    doc.add_paragraph('Analiza dzialania i logow\n', style='MyTitle').alignment = 1
    doc.add_paragraph('Komputerow Emisyjnych\n', style='Normal').alignment = 1

def add_introduction(doc, date):
    """
    Dodaje wstepny paragraf o dacie przegladu.
    """
    doc.add_paragraph(
        f'\tW dniu {date} roku przedstawiciel firmy '
        'dokonal przegladu Komputerow Emisyjnych. W trakcie realizacji zadania '
        'stwierdzono co nastepuje:', style='Normal'
    )

def add_log_analysis_section(doc, data):
    """
    Dodaje sekcje z analiza logow oraz podsekcje dotyczace restartow, synchronizacji czasu i logow systemowych.
    """
    doc.add_paragraph('\nAnaliza zdarzen i logow', style='NormalBold')

    # Podsekcja: restarty komputerow
    doc.add_paragraph('\na) restarty komputerow', style='Normal')
    for i in range(1, 4):
        doc.add_paragraph(
            f'SSMS{i}: {data.get(f"restart_ssms{i}_date")} ( przeglad komputerow, roczny )', 
            style='TabulatorStyl'
        )

    # Podsekcja: synchronizacja czasu
    doc.add_paragraph('\nb) synchronizacja czasu:', style='Normal')
    doc.add_paragraph(
        'Realizacja synchronizacji czasu odbywa sie z zegarem ntpd 192.168.110.250 podczas startu oraz w 30 minucie kazdej godziny.',
        style='TabulatorStyl'
    )
    doc.add_paragraph('Na obu serwerach czas jest zgodny i zsynchronizowany.', style='TabulatorStyl')

    # Podsekcja: analiza logow systemowych
    doc.add_paragraph('\nc) analiza logow systemowych', style='Normal')
    doc.add_paragraph('Duza ilosc zapytan protokolem ssh z komputera 108.208.108.8 bez proby logowania ( co 4 minuty ).', style='TabulatorStyl')
    doc.add_paragraph('W ciagu ostatniego tygodnia nie odnotowano zadnego niepozadanego logowania.', style='TabulatorStyl')

    # Podsekcja: analiza logow aplikacyjnych
    doc.add_paragraph('\nd) analiza logow aplikacyjnych\n', style='Normal')
    doc.add_paragraph('\t- SCADA:\n\t\tW ciagu ostatniego tygodnia nie odnotowano zadnego niepozadanego logowania.', style='Normal')
    doc.add_paragraph('\n\t- Baza danych POSTGRES 9.1\n\t\tDuza ilosc zapytan protokolem ssh z komputera 108.208.108.8 bez proby logowania\n\t\t ( co 4 minuty ).\n\t\tW ciagu ostatniego tygodnia nie odnotowano zadnego niepozadanego logowania.', style='Normal')

    # Podsekcja: dostepnosc portow tcp/udp
    doc.add_paragraph('\ne)  dostepnosc portow tcp/udp', style='Normal')
    doc.add_paragraph('\t22,  23 (sshd), 514, 6022, 6023, 6024, 6088, 5432, 502, 873', style='Normal')

def add_computer_load_section(doc, data):
    """
    Dodaje sekcje dotyczaca obciazen komputerow oraz zasobu dyskowego, w tym tabele z danymi dyskow.
    """
    doc.add_paragraph('\nf) obciazenia komputerow oraz zasobu dyskowego', style='Normal')
    for computer, specs in data["computers"].items():
        doc.add_paragraph(f'\n{computer}:')
        doc.add_paragraph(f'• srednie zuzycie procesora na poziomie {specs["cpu_usage"]}')
        doc.add_paragraph(f'• pamiec operacyjna RAM zajeta w {specs["ram_usage_percent"]} - zajete okolo {specs["ram_usage_mb"]} na 4GB dostepnych')
        doc.add_paragraph('• dysk twardy (ustawiony okres archiwizacji 2 lata)')

        add_disk_table(doc, specs["disk"])

def add_disk_table(doc, disk_data):
    """
    Dodaje tabele z danymi o zasobach dyskowych dla kazdego komputera.
    """
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Filesystem'
    hdr_cells[1].text = 'Size'
    hdr_cells[2].text = 'Used'
    hdr_cells[3].text = 'Available'
    hdr_cells[4].text = 'Capacity'
    hdr_cells[5].text = 'Mounted on'

    for disk in disk_data:
        row = table.add_row().cells
        row[0].text = disk['filesystem']
        row[1].text = disk['size']
        row[2].text = disk['used']
        row[3].text = disk['available']
        row[4].text = disk['capacity']
        row[5].text = disk['mounted_on']

def add_conclusions(doc):
    """
    Dodaje sekcje z wnioskami na koncu dokumentu.
    """
    doc.add_paragraph('\n\n\n\nWnioski\n', style='NormalBold')
    doc.add_paragraph(
        '\tNa podstawie zebranych danych mozna stwierdzic, ze pracujace w Elektrowni Komputery Emisyjne '
        ' dzialaja prawidlowo i nie zarejestrowano na nich zadnych prob niepozadanego uwierzytelniania.', 
        style='Normal'
    )

def generate_document(year, month, day, data):
    """
    Funkcja glowna do generowania dokumentu.
    """
    report_date = date(year, month, day)
    polish_month = get_polish_month(month)
    date_long = f'{day} {polish_month} {year}'

    doc = create_document()
    add_document_title(doc, report_date.strftime('%d.%m.%Y'))
    add_introduction(doc, date_long)
    add_log_analysis_section(doc, data)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_computer_load_section(doc, data)
    add_conclusions(doc)
    
    doc.save(f'analiza_logow_{year}_{month}_{day}.docx')
